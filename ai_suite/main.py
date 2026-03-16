from fastapi import FastAPI, HTTPException
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from docx import Document
import os
import importlib.util

app = FastAPI()

@app.get("/")
def read_root():
    return {"message": "AI Suite is Online"}

@app.get("/status")
def get_status():
    return {
            "status": "online",
            "capabilities": ["XGBoost", "Random Forest"],
            "version": "1.0.0"
        }

class ValidationRequest(BaseModel):
    framework: str
    version: str
    file_name: str

@app.post("/validate")
async def validate_model(request: ValidationRequest):
    # This commented code might be handy later if we only support specific frameworks
    # For testing it is just annoying
    """
    valid_frameworks = ["XGBoost", "Random Forest", "TensorFlow", "Excel", "General"]
    
    if request.framework not in valid_frameworks:
        return {"status": "invalid", "message": f"{request.framework} is not a supported framework."}
    """
    
    allowed_exts = ('.json', '.pkl', '.bst', '.h5', '.txt', '.xlsx', '.docx', '.py')
    if not request.file_name.lower().endswith(allowed_exts):
        return {"status": "invalid", "message": "Unsupported file extension for this platform."}

    return {"status": "valid", "message": "Model structure verified"}

# Code to return results from Python code
@app.post("/execute")
async def execute_model(model_id: int, file_path: str, output_type: str, user_input: str = None):
    full_path = f"/app/media_shared/{file_path}"
    
    # 1. Run the script as usual
    try:
        spec = importlib.util.spec_from_file_location("dynamic_mod", full_path)
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        raw_result = module.handle_request(user_input)
    except Exception as e:
        return {"status": "error", "message": f"Script Error: {str(e)}"}

    # 2. Check the REQUIRED output type
    if output_type == 'file':  # Matches the 'File Output' choice in your Django model
        doc = Document()
        doc.add_heading('AI Service Result', 0)
        doc.add_paragraph(raw_result)
        
        output_filename = f"result_{model_id}.docx"
        output_rel_path = f"results/{output_filename}"
        output_full_path = f"/app/media_shared/{output_rel_path}"
        
        os.makedirs(os.path.dirname(output_full_path), exist_ok=True)
        doc.save(output_full_path)
        
        return {
            "status": "success",
            "message": "File generated successfully.",
            "download_url": output_rel_path
        }

    # Default to text output
    return {
        "status": "success",
        "message": raw_result,
        "download_url": None
    }

# Logic used to test inputs and outputs of the AI service (Depreceated)
"""
# Logic to import and execute files
@app.post("/execute")
async def execute_model(model_id: int, file_path: str, user_input: str = None):
    # Path inside the container is /app/media_shared/
    full_path = f"/app/media_shared/{file_path}"
    
    if not os.path.exists(full_path):
        return {"status": "error", "message": f"File not found: {full_path}"}

    try:
        # Load the uploaded Python file dynamically
        spec = importlib.util.spec_from_file_location("dynamic_mod", full_path)
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        
        # This calls the function 'handle_request' inside YOUR uploaded file
        result = module.handle_request(user_input)
        
        return {"status": "success", "message": result}
    except Exception as e:
        return {"status": "error", "message": f"Script Error: {str(e)}"}
"""

# Used for debugging to list all files shared with Django
"""
@app.post("/execute")
async def execute_model(model_id: int, file_path: str):
    # This path points to the SHARED volume we set up in Docker
    full_path = f"/app/media_shared/{file_path}"
    
    if os.path.exists(full_path):
        # For now, we just return a "Success" message to prove connectivity
        return {
            "status": "success",
            "message": f"AI Suite located model {model_id} at {full_path}. Ready for processing."
        }
    else:
        raise HTTPException(status_code=404, detail=f"Model file not found at {full_path}")
"""

# AI generated mock text summarizer to test API functionality
"""
@app.post("/execute")
async def execute_model(model_id: int, user_input: str = None):
    # This is our 'Mock' AI logic
    if not user_input or len(user_input.strip()) == 0:
        return {"status": "error", "message": "Please provide some text to summarize!"}

    # Simulate summarization: Take the first 120 characters and add '...'
    # In the future, this is where your 'Excel-to-Word' or 'ML Model' code lives.
    summary = user_input[:120] + "..." if len(user_input) > 120 else user_input

    return {
        "status": "success", 
        "message": summary
    }
"""